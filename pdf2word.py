#!/usr/bin/env python3

"""
Put each page of a PDF file (from a scanner) as an image into a WORD document

Requirements:
pip install -r requirements.txt

Usage:
1 - download poppler from https://poppler.freedesktop.org/
2 - save to the same folder, script uses ./poppler/bin as default path
3 - run, select a PDF file => word docx
"""

import docx
from docx.shared import Inches
from tkinter import filedialog
from tkinter import Tk
from pdf2image import convert_from_path
from PyPDF2 import PdfFileReader, PdfFileWriter
import os


def pdf2word():
    root = Tk()
    root.withdraw()

    my_file = filedialog.askopenfilename()
    print('File: {}'.format(my_file))
    if my_file == '':
        print('No file selected')
        return 

    pdf_rd = PdfFileReader(my_file)
    doc = docx.Document()

    print('Total Pages: {}'.format(pdf_rd.getNumPages()))
    for i in range(pdf_rd.getNumPages()):
        print('Processing Page: {:03}'.format(i+1))
        page = pdf_rd.getPage(i)
        pdf_wr = PdfFileWriter()
        pdf_wr.addPage(page)
        tmp_pdf_file_name = 'tmp_pdf.pdf'
        f = open(tmp_pdf_file_name, 'wb')
        pdf_wr.write(f)
        f.close()
        images = convert_from_path(tmp_pdf_file_name,
                                   dpi=300,
                                   poppler_path='./poppler/bin')
        os.remove(tmp_pdf_file_name)

        ratio = 0.68
        for img in images:
            f_name = 'tmp.jpg'
            img.save(f_name)
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_break()
            doc.add_picture(f_name,
                            width=Inches(ratio * 8.5))
            doc.add_page_break()
            os.remove(f_name)


    doc_fname = os.path.splitext(my_file)[0] + '.docx'
    print('Saving to {}'.format(doc_fname))
    doc.save(doc_fname)



if __name__ == '__main__':
    pdf2word()